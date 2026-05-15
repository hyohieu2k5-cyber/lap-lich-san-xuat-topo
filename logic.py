from datetime import datetime
import pandas as pd
from docx import Document
from tkinter import filedialog
from tkinter import messagebox, ttk

# logic.py
import pandas as pd
from tkinter import filedialog

# Quản lý Frames
frames = {}

def register_frames(h_frame, w_frame):
    frames['home'] = h_frame
    frames['working'] = w_frame

def switch_to_home():
    if 'working' in frames: frames['working'].pack_forget()
    if 'home' in frames: frames['home'].pack(expand=True, fill="both")

def switch_to_scheduler():
    if 'home' in frames: frames['home'].pack_forget()
    if 'working' in frames: frames['working'].pack(expand=True, fill="both")

# Logic chuẩn hóa dữ liệu nhập khẩu
def import_from_file():
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx *.xls")])
    if not file_path: return None
    
    try:
        df = pd.read_excel(file_path)
        raw_data = df.to_dict(orient='records')
        standardized = []
        
        for item in raw_data:
            # Tìm ID (chấp nhận id, ID, Id)
            t_id = next((v for k, v in item.items() if k.lower() == 'id'), None)
            # Tìm Name
            t_name = next((v for k, v in item.items() if k.lower() in ['name', 'tên công đoạn', 'tên']), None)
            # Tìm Duration
            t_dur = next((v for k, v in item.items() if k.lower() in ['duration', 'thời gian', 'thời gian (h)']), 0)
            # Tìm Date
            t_date = next((v for k, v in item.items() if k.lower() in ['date', 'ngày', 'ngày thực hiện']), "01/01/2026")
            if isinstance(t_date, datetime):
                t_date = t_date.strftime("%d/%m/%Y")
            else:
                t_date = str(t_date)
            # Tìm Priority
            t_priority = next((v for k, v in item.items() if k and k.strip().lower() in ['priority', 'ưu tiên']), 0)
            standardized.append({
                    "id": int(t_id),
                    "name": str(t_name),
                    "duration": int(t_dur),
                    "date": str(t_date),
                    "priority": int(t_priority)
                })
        return standardized
    except Exception as e:
        print(f"Lỗi: {e}")
        return None
    
from datetime import datetime

def calculate_topo(tasks, edges):
    if not tasks:
        return "Trống"

    in_degree = {t['id']: 0 for t in tasks}
    adj = {t['id']: [] for t in tasks}

    for u, v in edges:
        if u in adj and v in adj:
            adj[u].append(v)
            in_degree[v] += 1

    # 🔥 danh sách sẵn sàng
    ready = []
    for t in tasks:
        if in_degree[t['id']] == 0:
            date_obj = datetime.strptime(t['date'], "%d/%m/%Y")
            priority = t.get("priority", 0)

            # 🔥 ưu tiên cao trước → dùng -priority
            ready.append((-priority, date_obj, t['id']))

    ready.sort(key=lambda x: (x[0], x[1]))

    task_map = {t['id']: t for t in tasks}

    result = []

    while ready:
        _, _, u = ready.pop(0)
        result.append(u)

        for v in adj[u]:
            in_degree[v] -= 1
            if in_degree[v] == 0:
                v_task = task_map[v]
                v_date = datetime.strptime(v_task['date'], "%d/%m/%Y")
                v_priority = v_task.get("priority", 0)

                ready.append((-v_priority, v_date, v))
                ready.sort(key=lambda x: (x[0], x[1]))

    if len(result) != len(tasks):
        return None

    task_map = {t['id']: t for t in tasks}

    return "\n".join([
    f"{i+1}. [{task_map[t]['date']}] (P{task_map[t].get('priority',0)}) {task_map[t]['name']}"
    for i, t in enumerate(result)
])

# Thêm 2 hàm này để fix lỗi ImportError ở main.py
def save_data(tasks, edges):
    pass

def load_data():
    return [], []

#---CẬP NHẬT---#

def update_task(tasks, t_id, new_name, new_duration, new_date, new_priority):
    for t in tasks:
        if t["id"] == t_id:
            t["name"] = new_name
            t["duration"] = new_duration
            t["date"] = new_date
            t["priority"] = new_priority
            return True
    return False

#---XÓA---#
def delete_task(tasks, t_id):
    new_tasks = [t for t in tasks if t["id"] != t_id]
    return new_tasks

# --- XUẤT EXCEL ---
def export_to_excel(tasks):
    if not tasks: return False
    df = pd.DataFrame(tasks)
    # Đổi tên cột cho đẹp
    df.columns = ['ID', 'Tên công đoạn', 'Thời gian (h)', 'Ngày thực hiện', "Ưu tiên"]
    file_path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
    if file_path:
        df.to_excel(file_path, index=False)
        return True
    return False

# --- XUẤT WORD ---
def export_to_word(tasks, result_topo):
    if not tasks: return False
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
    if file_path:
        doc = Document()
        doc.add_heading('BÁO CÁO KẾ HOẠCH SẢN XUẤT', 0)
        
        # Thêm bảng dữ liệu
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'; hdr_cells[1].text = 'Tên'; hdr_cells[2].text = 'Thời gian'; hdr_cells[3].text = 'Ngày'; hdr_cells[4].text = 'Ưu tiên'
        
        for t in tasks:
            row = table.add_row().cells
            row[0].text = str(t['id'])
            row[1].text = t['name']
            row[2].text = f"{t['duration']}h"
            row[3].text = t['date']
            row[4].text = str(t.get('priority', 0))
            
        doc.add_heading('Thứ tự thực hiện (Topo):', level=1)
        doc.add_paragraph(result_topo if result_topo else "Chưa lập lịch")
        
        doc.save(file_path)
        return True
    return False

#---HƯỚNG DẪN SỬ DỤNG---#
def show_guide():
    messagebox.showinfo("Hướng dẫn sử dụng",
        "📌 CÁCH DÙNG PHẦN MỀM:\n\n"
        "1. Nhập ID, Tên, Thời gian, chọn ngày\n"
        "2. Nhấn 'Thêm công đoạn'\n"
        "3. Thêm quan hệ trước - sau (ID trước -> ID sau)\n"
        "4. Nhấn 'Chạy lập lịch' để xem thứ tự\n\n"
        "📁 Có thể xuất Excel/Word hoặc nhập dữ liệu\n"
        "⚠️ Không được tạo chu trình (vòng lặp)"
    )